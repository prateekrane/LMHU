from docx import Document
import logging

class WordBuilder:
    def __init__(self):
        self.doc = Document()

    def run_command(self, command):
        action = command.get("action")
        content = command.get("content", "")
        try:
            if action == "add_heading":
                self.doc.add_heading(content, level=1)
            elif action == "add_paragraph":
                self.doc.add_paragraph(content)
            elif action == "add_bullet":
                self.doc.add_paragraph(content, style='ListBullet')
            else:
                logging.warning(f"Unknown Word command: {action}")
        except Exception as e:
            logging.error(f"Failed to run Word command: {command}", exc_info=True)

    def save(self, filepath):
        self.doc.save(filepath)