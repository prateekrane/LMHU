from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import json

# Supported Document Options (pass in data['formatting'] or data['style_settings']):
# - page_size: 'a4', 'letter', 'legal', 'a3'
# - orientation: 'portrait', 'landscape'
# - margins: {'top': float, 'bottom': float, 'left': float, 'right': float} (in inches)
# - include_page_numbers: bool
# - page_number_text: str
# - page_number_of_total: bool
# - include_header: bool
# - header_text: str
# - header_include_date: bool
# - include_footer: bool
# - footer_text: str
# - footer_include_date: bool
# - text_alignment: {title, headings, body}
# - styles: {custom styles}
# - default_paragraph: {style settings}
# - headings: {level: {style settings}}
# - title_style, default_run: {style settings}
# - ...
#
# Example usage:
# data = {
#   'formatting': {
#       'page_size': 'a4',
#       'orientation': 'landscape',
#       'margins': {'top': 1, 'bottom': 1, 'left': 1, 'right': 1},
#       'include_page_numbers': True,
#       ...
#   },
#   ...
# }

def _clean_text(text):
    """Clean text by removing invalid escape characters and normalizing whitespace."""
    if not text:
        return ""
    # Replace common escape sequences with their proper characters
    text = text.replace("\\n", "\n")
    text = text.replace("\\t", "\t")
    text = text.replace("\\r", "\r")
    # Remove any remaining invalid escape sequences
    text = text.encode('utf-8', 'ignore').decode('utf-8')
    
    # Also remove markdown characters like *, **, and leading bullets
    import re
    text = re.sub(r'^[\s\*\-]+', '', text) # Remove leading *, -, or numbers with dot
    text = text.replace('**', '').replace('*', '') # Remove asterisks
    
    return text.strip()

def _set_element_style(element, style_settings):
    """Apply style settings to a paragraph or run."""
    if not style_settings:
        return

    if hasattr(element, 'paragraph_format'):
        # Paragraph formatting
        if 'alignment' in style_settings:
            try:
                alignment = style_settings['alignment'].upper()
                element.paragraph_format.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment, WD_ALIGN_PARAGRAPH.LEFT)
            except (AttributeError, ValueError):
                pass
        if 'line_spacing' in style_settings:
            try:
                element.paragraph_format.line_spacing = style_settings['line_spacing']
            except (AttributeError, ValueError):
                pass
        if 'space_before' in style_settings:
            try:
                element.paragraph_format.space_before = Pt(style_settings['space_before'])
            except (AttributeError, ValueError):
                pass
        if 'space_after' in style_settings:
            try:
                element.paragraph_format.space_after = Pt(style_settings['space_after'])
            except (AttributeError, ValueError):
                pass
        if 'indentation' in style_settings:
            try:
                element.paragraph_format.left_indent = Pt(style_settings['indentation'])
            except (AttributeError, ValueError):
                pass
    
    if hasattr(element, 'font'):
        # Font formatting
        if 'name' in style_settings:
            try:
                element.font.name = style_settings['name']
                if hasattr(element, '_element'):
                    # Create or get rPr element
                    if not hasattr(element._element, 'rPr') or element._element.rPr is None:
                        rPr = OxmlElement('w:rPr')
                        element._element.append(rPr)
                    else:
                        rPr = element._element.rPr
                    
                    # Create or get rFonts element
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:eastAsia'), style_settings['name'])
                    rFonts.set(qn('w:ascii'), style_settings['name'])
                    rFonts.set(qn('w:hAnsi'), style_settings['name'])
                    
                    # Remove existing rFonts if any
                    existing_rFonts = rPr.find(qn('w:rFonts'))
                    if existing_rFonts is not None:
                        rPr.remove(existing_rFonts)
                    
                    # Add the new rFonts
                    rPr.append(rFonts)
            except (AttributeError, ValueError):
                pass
        if 'size' in style_settings:
            try:
                element.font.size = Pt(style_settings['size'])
            except (AttributeError, ValueError):
                pass
        if 'bold' in style_settings:
            try:
                element.font.bold = style_settings['bold']
            except (AttributeError, ValueError):
                pass
        if 'italic' in style_settings:
            try:
                element.font.italic = style_settings['italic']
            except (AttributeError, ValueError):
                pass
        if 'underline' in style_settings:
            try:
                element.font.underline = style_settings['underline']
            except (AttributeError, ValueError):
                pass
        if 'color' in style_settings:
            try:
                color = style_settings['color']
                if isinstance(color, str) and color.startswith('#'):
                    color = color[1:]  # Remove # if present
                    r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
                    element.font.color.rgb = RGBColor(r, g, b)
            except (AttributeError, ValueError):
                pass

def _create_style(doc, style_name, style_settings):
    """Create a custom style in the document."""
    try:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        _set_element_style(style, style_settings)
        return style
    except ValueError:
        # Style already exists
        return doc.styles[style_name]

def _apply_document_style(doc, style_settings):
    """Apply document-wide style settings."""
    if not style_settings:
        return

    # First, ensure we have a base style to inherit from
    base_style = doc.styles['Normal']
    
    # Apply base font and alignment to the document's default style
    if 'font' in style_settings:
        base_style.font.name = style_settings['font']
        if hasattr(base_style, '_element'):
            # Create or get rPr element
            if not hasattr(base_style._element, 'rPr') or base_style._element.rPr is None:
                rPr = OxmlElement('w:rPr')
                base_style._element.append(rPr)
            else:
                rPr = base_style._element.rPr
            
            # Create or get rFonts element
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), style_settings['font'])
            rFonts.set(qn('w:ascii'), style_settings['font'])
            rFonts.set(qn('w:hAnsi'), style_settings['font'])
            
            # Remove existing rFonts if any
            existing_rFonts = rPr.find(qn('w:rFonts'))
            if existing_rFonts is not None:
                rPr.remove(existing_rFonts)
            
            # Add the new rFonts
            rPr.append(rFonts)
    
    if 'text_alignment' in style_settings:
        alignment = style_settings['text_alignment'].upper()
        base_style.paragraph_format.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Apply default paragraph style with all settings
    if 'default_paragraph' in style_settings:
        para_settings = style_settings['default_paragraph']
        # Ensure font is set from document style if not specified
        if 'font' in style_settings and 'name' not in para_settings:
            para_settings['name'] = style_settings['font']
        # Ensure alignment is set from document style if not specified
        if 'text_alignment' in style_settings and 'alignment' not in para_settings:
            para_settings['alignment'] = style_settings['text_alignment']
        _set_element_style(base_style, para_settings)

    # Apply heading styles with proper inheritance
    if 'headings' in style_settings:
        for level, settings in style_settings['headings'].items():
            style_name = f'Heading {level}'
            if style_name in doc.styles:
                style = doc.styles[style_name]
                # Ensure font is set from document style if not specified
                if 'font' in style_settings and 'name' not in settings:
                    settings['name'] = style_settings['font']
                # Ensure alignment is set from document style if not specified
                if 'text_alignment' in style_settings and 'alignment' not in settings:
                    settings['alignment'] = style_settings['text_alignment']
                _set_element_style(style, settings)

    # Apply title style
    if 'title_style' in style_settings:
        title_settings = style_settings['title_style']
        if 'Title' in doc.styles:
            style = doc.styles['Title']
            # Ensure font is set from document style if not specified
            if 'font' in style_settings and 'name' not in title_settings:
                title_settings['name'] = style_settings['font']
            # Ensure alignment is set from document style if not specified
            if 'text_alignment' in style_settings and 'alignment' not in title_settings:
                title_settings['alignment'] = style_settings['text_alignment']
            _set_element_style(style, title_settings)

    # Apply to all other styles that should inherit document settings
    all_styles = [
        'Normal', 'Title', 'TOC 1', 'TOC 2', 'TOC 3', 
        'Header', 'Footer', 'Default Paragraph Font',
        'List Paragraph', 'Quote', 'Intense Quote'
    ]
    all_styles.extend([f'Heading {i}' for i in range(1, 10)])

    for style_name in all_styles:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            
            # Apply font if specified in document style
            if 'font' in style_settings:
                try:
                    style.font.name = style_settings['font']
                    if hasattr(style, '_element'):
                        # Create or get rPr element
                        if not hasattr(style._element, 'rPr') or style._element.rPr is None:
                            rPr = OxmlElement('w:rPr')
                            style._element.append(rPr)
                        else:
                            rPr = style._element.rPr
                        
                        # Create or get rFonts element
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:eastAsia'), style_settings['font'])
                        rFonts.set(qn('w:ascii'), style_settings['font'])
                        rFonts.set(qn('w:hAnsi'), style_settings['font'])
                        
                        # Remove existing rFonts if any
                        existing_rFonts = rPr.find(qn('w:rFonts'))
                        if existing_rFonts is not None:
                            rPr.remove(existing_rFonts)
                        
                        # Add the new rFonts
                        rPr.append(rFonts)
                except (AttributeError, ValueError):
                    pass

            # Apply alignment if specified in document style
            if 'text_alignment' in style_settings:
                try:
                    alignment = style_settings['text_alignment'].upper()
                    style.paragraph_format.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment, WD_ALIGN_PARAGRAPH.LEFT)
                except (AttributeError, ValueError):
                    pass

    # Apply custom styles if specified
    if 'styles' in style_settings:
        for style_name, settings in style_settings['styles'].items():
            try:
                style = _create_style(doc, style_name, settings)
                # Ensure font and alignment are inherited from document style if not specified
                if 'font' in style_settings and 'name' not in settings:
                    style.font.name = style_settings['font']
                if 'text_alignment' in style_settings and 'alignment' not in settings:
                    style.paragraph_format.alignment = getattr(WD_ALIGN_PARAGRAPH, style_settings['text_alignment'].upper(), WD_ALIGN_PARAGRAPH.LEFT)
            except Exception:
                continue

def _add_page_number_field(paragraph, text_before="", text_after=""):
    """Add a page number field to a paragraph using python-docx API."""
    if text_before:
        paragraph.add_run(_clean_text(text_before))

    # Add the page number field using a field code
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    if text_after:
        paragraph.add_run(_clean_text(text_after))

def _add_total_pages_field(paragraph, text_before="", text_after=""):
    """Add a total pages field to a paragraph using python-docx API."""
    if text_before:
        paragraph.add_run(_clean_text(text_before))

    # Add the total pages field using a field code
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def _clean_toc_and_md(text):
    """Remove asterisks, double asterisks, and leading markdown bullets from text."""
    if not text:
        return ""
    # Remove leading * or - or numbers with dot (markdown bullets)
    import re
    text = re.sub(r'^[\s\*\-]+', '', text)
    text = text.replace('**', '').replace('*', '')
    return text.strip()

def _apply_formatting(doc, formatting):
    """Apply document formatting settings (only options from UI)."""
    if not formatting:
        return

    # Remove TOC feature
    if "include_toc" in formatting:
        del formatting["include_toc"]

    # Set text alignment for different elements
    if "text_alignment" in formatting:
        alignment = formatting["text_alignment"].upper()
        for style_name in ["Title", "Normal", "TOC 1"]:
            if style_name in doc.styles:
                doc.styles[style_name].paragraph_format.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment, WD_ALIGN_PARAGRAPH.LEFT)
        for i in range(1, 10):
            style_name = f"Heading {i}"
            if style_name in doc.styles:
                doc.styles[style_name].paragraph_format.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Set font for all major styles (restricted to allowed fonts)
    allowed_fonts = ["Times New Roman", "Calibri", "Arial", "Georgia", "Verdana"]
    if "font" in formatting and formatting["font"]:
        font_name = formatting["font"]
        if font_name not in allowed_fonts:
            font_name = "Calibri"  # Default fallback
        for style_name in ["Normal", "Title", "TOC 1"]:
            if style_name in doc.styles:
                doc.styles[style_name].font.name = font_name
                doc.styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        for i in range(1, 10):
            style_name = f"Heading {i}"
            if style_name in doc.styles:
                doc.styles[style_name].font.name = font_name
                doc.styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Set page size
    if "page_size" in formatting:
        page_size = formatting["page_size"].lower()
        if page_size == "a4":
            doc.sections[0].page_width = Inches(8.27)
            doc.sections[0].page_height = Inches(11.69)
        elif page_size == "letter":
            doc.sections[0].page_width = Inches(8.5)
            doc.sections[0].page_height = Inches(11)
        elif page_size == "legal":
            doc.sections[0].page_width = Inches(8.5)
            doc.sections[0].page_height = Inches(14)
        elif page_size == "a3":
            doc.sections[0].page_width = Inches(11.69)
            doc.sections[0].page_height = Inches(16.54)
        elif page_size == "a5":
            doc.sections[0].page_width = Inches(5.83)
            doc.sections[0].page_height = Inches(8.27)

    # Add page numbers
    if formatting.get("include_page_numbers", False):
        section = doc.sections[0]
        footer = section.footer
        # Clear all existing paragraphs in the footer
        for para in list(footer.paragraphs):
            p = para._element
            p.getparent().remove(p)
        # Always create a new paragraph for the page number
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_text = formatting.get("page_number_text", "Page ")
        if formatting.get("page_number_of_total", False):
            _add_page_number_field(paragraph, text_before=page_text)
            paragraph.add_run(" of ")
            _add_total_pages_field(paragraph)
        else:
            _add_page_number_field(paragraph, text_before=page_text)

    # Add header if selected
    if formatting.get("include_header", False):
        section = doc.sections[0]
        header = section.header
        # Clear all existing paragraphs in the header
        for para in list(header.paragraphs):
            p = para._element
            p.getparent().remove(p)
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_text = formatting.get("header_text", doc.core_properties.title or "Document")
        if formatting.get("header_include_date", False):
            header_text += f" - {datetime.now().strftime('%Y-%m-%d')}"
        paragraph.text = _clean_text(header_text)

    # Add footer if selected
    if formatting.get("include_footer", False):
        section = doc.sections[0]
        footer = section.footer
        # Add a new paragraph for footer text, but do not remove the page number paragraph
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = formatting.get("footer_text", "")
        if formatting.get("footer_include_date", False):
            if footer_text:
                footer_text += " - "
            footer_text += datetime.now().strftime('%Y-%m-%d')
        paragraph.text = _clean_text(footer_text)

    # Clean all headings and content
    for section_data in formatting.get("sections", []):
        section_data["heading"] = _clean_text(section_data.get("heading", ""))
        section_data["content"] = _clean_text(section_data.get("content", ""))

def get_document_style_preset(style_name):
    """Return style settings for a given document style preset."""
    presets = {
        "professional": {
            "font": "Calibri",
            "text_alignment": "JUSTIFY",
            "default_paragraph": {
                "name": "Calibri",
                "size": 12,
                "color": "#222222",
                "line_spacing": 1.15,
                "space_before": 6,
                "space_after": 6,
            },
            "headings": {
                "1": {"name": "Calibri", "size": 16, "bold": True, "color": "#2E74B5"},
                "2": {"name": "Calibri", "size": 14, "bold": True, "color": "#5B9BD5"},
            },
            "title_style": {"name": "Calibri", "size": 22, "bold": True, "color": "#1F4E79"},
            "default_run": {"name": "Calibri", "size": 12},
        },
        "business": {
            "font": "Arial",
            "text_alignment": "LEFT",
            "default_paragraph": {
                "name": "Arial",
                "size": 11,
                "color": "#222222",
                "line_spacing": 1.2,
                "space_before": 4,
                "space_after": 4,
            },
            "headings": {
                "1": {"name": "Arial", "size": 16, "bold": True, "color": "#0070C0"},
                "2": {"name": "Arial", "size": 13, "bold": True, "color": "#0070C0"},
            },
            "title_style": {"name": "Arial", "size": 20, "bold": True, "color": "#0070C0"},
            "default_run": {"name": "Arial", "size": 11},
        },
        "academic": {
            "font": "Times New Roman",
            "text_alignment": "JUSTIFY",
            "default_paragraph": {
                "name": "Times New Roman",
                "size": 12,
                "color": "#222222",
                "line_spacing": 2.0,
                "space_before": 0,
                "space_after": 8,
            },
            "headings": {
                "1": {"name": "Times New Roman", "size": 16, "bold": True, "color": "#000000"},
                "2": {"name": "Times New Roman", "size": 14, "bold": True, "color": "#000000"},
            },
            "title_style": {"name": "Times New Roman", "size": 22, "bold": True, "color": "#000000"},
            "default_run": {"name": "Times New Roman", "size": 12},
        },
        "creative": {
            "font": "Georgia",
            "text_alignment": "LEFT",
            "default_paragraph": {
                "name": "Georgia",
                "size": 12,
                "color": "#333333",
                "line_spacing": 1.3,
                "space_before": 8,
                "space_after": 8,
            },
            "headings": {
                "1": {"name": "Georgia", "size": 18, "bold": True, "color": "#C00000"},
                "2": {"name": "Georgia", "size": 15, "bold": True, "color": "#FF6600"},
            },
            "title_style": {"name": "Georgia", "size": 24, "bold": True, "color": "#C00000"},
            "default_run": {"name": "Georgia", "size": 12},
        },
        "technical": {
            "font": "Verdana",
            "text_alignment": "LEFT",
            "default_paragraph": {
                "name": "Verdana",
                "size": 11,
                "color": "#222222",
                "line_spacing": 1.1,
                "space_before": 2,
                "space_after": 6,
            },
            "headings": {
                "1": {"name": "Verdana", "size": 15, "bold": True, "color": "#1E4E79"},
                "2": {"name": "Verdana", "size": 13, "bold": True, "color": "#2E75B6"},
            },
            "title_style": {"name": "Verdana", "size": 20, "bold": True, "color": "#1E4E79"},
            "default_run": {"name": "Verdana", "size": 11},
        },
    }
    return presets.get(style_name.lower(), {})

def generate_word_doc(data):
    try:
        # Create document
        doc = Document()
        
        # Set document properties
        title = _clean_text(data.get("title", "Untitled Document"))
        doc.core_properties.title = title
        if "author" in data:
            doc.core_properties.author = data["author"]
        if "subject" in data:
            doc.core_properties.subject = data["subject"]
        if "keywords" in data:
            doc.core_properties.keywords = data["keywords"]
        
        # --- FIX: Use 'style' key for Document Style preset ---
        style_name = data.get("style") or data.get("formatting", {}).get("style")
        style_settings = data.get("style_settings", {}) or {}
        if style_name:
            preset = get_document_style_preset(style_name)
            import copy
            def deep_merge(a, b):
                for k, v in b.items():
                    if k in a and isinstance(a[k], dict) and isinstance(v, dict):
                        deep_merge(a[k], v)
                    else:
                        a[k] = v
                return a
            style_settings = deep_merge(copy.deepcopy(style_settings), preset)

        # --- OVERRIDE: Font and Text Alignment if selected ---
        formatting = data.get("formatting", {})
        font_override = formatting.get("font", None)
        align_override = formatting.get("text_alignment", None)
        if font_override and font_override.lower() != "none":
            style_settings["font"] = font_override
            if "default_paragraph" in style_settings:
                style_settings["default_paragraph"]["name"] = font_override
            if "default_run" in style_settings:
                style_settings["default_run"]["name"] = font_override
            if "title_style" in style_settings:
                style_settings["title_style"]["name"] = font_override
            if "headings" in style_settings:
                for k in style_settings["headings"]:
                    style_settings["headings"][k]["name"] = font_override
        if align_override and align_override.lower() != "none":
            style_settings["text_alignment"] = align_override
        
        # Apply document style
        _apply_document_style(doc, style_settings)
        
        # Apply formatting
        _apply_formatting(doc, formatting)
        
        # Add title with custom style
        title_style = style_settings.get("title_style", {})
        title_paragraph = doc.add_heading(title, level=0)
        _set_element_style(title_paragraph, title_style)
        
        # Add sections
        for section in data.get("sections", []):
            heading = _clean_text(section.get("heading", ""))
            content = _clean_text(section.get("content", ""))
            level = section.get("level", 1)
            
            if heading:
                # Apply heading style
                heading_style = style_settings.get("headings", {}).get(str(level), {})
                heading_paragraph = doc.add_heading(heading, level=level)
                _set_element_style(heading_paragraph, heading_style)
            
            if content:
                # Split content into paragraphs and add them
                paragraphs = content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Apply paragraph style
                        paragraph_style = style_settings.get("default_paragraph", {})
                        p = doc.add_paragraph()
                        _set_element_style(p, paragraph_style)
                        
                        # Add text with run style
                        run_style = style_settings.get("default_run", {})
                        run = p.add_run(para_text.strip())
                        _set_element_style(run, run_style)
        
        # Generate filename and save path
        downloads_path = os.path.expanduser("~/Downloads")
        os.makedirs(downloads_path, exist_ok=True)
        
        clean_title = "".join(c for c in title if c.isalnum() or c.isspace())
        clean_title = clean_title.replace(" ", "_")[:30]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{clean_title}_{timestamp}.docx"
        filepath = os.path.join(downloads_path, filename)
        
        # Save document using a context manager to ensure proper closing
        with open(filepath, 'wb') as f:
            doc.save(f)
            
        # Explicitly delete the document object after saving
        del doc
        
        return filepath
        
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        raise
